"""
POST /api/v1/extract-text

Accepts an uploaded transcript file and returns the extracted plain text.

Supported formats:
    .txt  .md  .vtt  .srt    — read as UTF-8 / latin-1 text
    .docx                    — python-docx (Open Office XML)
    .pdf                     — pypdf

Legacy binary Word (.doc) is NOT supported — it requires heavyweight
extractors (antiword, libreoffice).  Users should save-as .docx first.

Size limit: 20 MB.  Larger uploads are rejected.

Security notes
--------------
* The filename is only used to pick an extractor by extension; it is never
  used to write to disk.  File content is processed in-memory.
* PDFs are parsed with pypdf's built-in text extractor, which does not
  execute embedded JavaScript or follow external references.
* .docx is processed via python-docx which reads the OOXML parts it knows
  about; we do not evaluate any embedded links or macros.
* The response returns only extracted text — no metadata about the
  uploading user, original filename, or server paths.
"""

import io
import logging

from fastapi import APIRouter, HTTPException, UploadFile, File

router = APIRouter()
logger = logging.getLogger(__name__)

# Maximum accepted upload size (bytes).  Keeps memory footprint bounded
# and prevents trivial DoS via giant PDFs.
_MAX_UPLOAD_BYTES = 20 * 1024 * 1024   # 20 MB

# Extensions that are plain text; just decode.
_TEXT_EXTS = {".txt", ".md", ".vtt", ".srt"}


def _ext_of(filename: str) -> str:
    """Return the lowercased extension including the dot, or ''."""
    name = (filename or "").lower().strip()
    idx = name.rfind(".")
    return name[idx:] if idx >= 0 else ""


def _decode_text(raw: bytes) -> str:
    """
    Decode bytes as UTF-8; fall back to latin-1 for legacy files that
    claim to be text but aren't valid UTF-8.  Neither raises.
    """
    try:
        return raw.decode("utf-8")
    except UnicodeDecodeError:
        return raw.decode("latin-1", errors="replace")


def _extract_docx(raw: bytes) -> str:
    """Extract paragraph text from a .docx file."""
    try:
        import docx  # python-docx
    except ImportError:
        raise HTTPException(
            status_code=503,
            detail="Server is missing 'python-docx'. Ask admin to pip install python-docx.",
        )
    try:
        doc = docx.Document(io.BytesIO(raw))
    except Exception as exc:
        raise HTTPException(
            status_code=400,
            detail=f"Could not read .docx file (is it corrupted or password-protected?): {exc}",
        )
    paragraphs = [p.text for p in doc.paragraphs if p.text]
    # Also pull table cell text — meeting-notes templates often use tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    paragraphs.append(cell.text)
    return "\n".join(paragraphs)


def _extract_pdf(raw: bytes) -> str:
    """Extract text content page-by-page from a .pdf file."""
    try:
        import pypdf
    except ImportError:
        raise HTTPException(
            status_code=503,
            detail="Server is missing 'pypdf'. Ask admin to pip install pypdf.",
        )
    try:
        reader = pypdf.PdfReader(io.BytesIO(raw))
    except Exception as exc:
        raise HTTPException(
            status_code=400,
            detail=f"Could not read PDF (is it corrupted or password-protected?): {exc}",
        )
    if reader.is_encrypted:
        # Attempt empty-password unlock; many "encrypted" PDFs have no password
        try:
            reader.decrypt("")
        except Exception:
            pass
        if reader.is_encrypted:
            raise HTTPException(
                status_code=400,
                detail="PDF is password-protected. Please unlock or export as .txt first.",
            )

    pages: list[str] = []
    for i, page in enumerate(reader.pages):
        try:
            pages.append(page.extract_text() or "")
        except Exception as exc:
            # Bad page shouldn't abort the whole file
            logger.warning("PDF page %d extraction failed: %s", i, exc)
    return "\n\n".join(p for p in pages if p.strip())


@router.post("/extract-text")
async def extract_text(file: UploadFile = File(...)) -> dict:
    """
    Accept a single file upload and return its extracted plain-text content.

    Request: multipart/form-data with a single ``file`` field.
    Response: ``{"text": "...", "format": ".docx", "bytes": 12345}``

    Errors:
      400 — unsupported format, corrupted file, or file too large
      503 — server missing an extractor dependency
    """
    # ``UploadFile.read()`` returns all bytes; hold our own size check since
    # UploadFile itself does not enforce a limit.
    raw = await file.read()
    if len(raw) > _MAX_UPLOAD_BYTES:
        raise HTTPException(
            status_code=400,
            detail=f"File too large ({len(raw):,} bytes). Maximum is {_MAX_UPLOAD_BYTES:,} bytes (20 MB).",
        )
    if not raw:
        raise HTTPException(status_code=400, detail="Empty file.")

    ext = _ext_of(file.filename or "")

    if ext in _TEXT_EXTS:
        text = _decode_text(raw)
    elif ext == ".docx":
        text = _extract_docx(raw)
    elif ext == ".pdf":
        text = _extract_pdf(raw)
    elif ext == ".doc":
        raise HTTPException(
            status_code=400,
            detail=(
                "Legacy .doc (Word 97-2003) is not supported. Please open the file "
                "in Word and Save As .docx, then upload again."
            ),
        )
    else:
        raise HTTPException(
            status_code=400,
            detail=(
                f"Unsupported file type: '{ext or '(no extension)'}'. "
                "Supported: .txt, .md, .vtt, .srt, .docx, .pdf"
            ),
        )

    # Strip leading/trailing blank lines but keep internal structure
    text = text.strip("\n\r ")

    if not text:
        raise HTTPException(
            status_code=400,
            detail=f"No readable text found in the uploaded {ext} file.",
        )

    return {
        "text":   text,
        "format": ext,
        "bytes":  len(raw),
        "chars":  len(text),
    }
